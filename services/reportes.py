from io import BytesIO
from datetime import datetime
import re
import csv
import io
import unicodedata
import os

from openpyxl import Workbook
from openpyxl.styles import Alignment, PatternFill, Font

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ══════════════════════════════════════════════════════════════════
# HELPERS XML (internos, no exponer fuera del módulo)
# ══════════════════════════════════════════════════════════════════
def _set_row_height(row, height_cm: float):
    """Fija la altura mínima de una fila en cm."""
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    for old in trPr.findall(qn('w:trHeight')):
        trPr.remove(old)
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'),   str(int(height_cm * 567)))
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)
def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Convierte '#1E1B4B' o '1E1B4B' a RGBColor."""
    h = hex_color.strip().lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_color: str):
    """Relleno sólido de celda."""
    h    = hex_color.strip().lstrip('#')
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # eliminar shd previo si existe
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  h.upper())
    tcPr.append(shd)


def _set_cell_borders(cell, hex_color: str = 'CCCCCC', size: str = '4'):
    """Bordes finos en todas las caras de la celda."""
    h    = hex_color.strip().lstrip('#').upper()
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    size)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), h)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_margin(cell, top=80, bottom=80, left=120, right=120):
    """Padding interno de celda en twips."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    mar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def _add_hr(doc, hex_color: str = '00B4D8', size: int = 10):
    """Línea horizontal decorativa debajo de un párrafo."""
    h = hex_color.strip().lstrip('#').upper()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(size))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), h)
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def _page_break(doc):
    p  = doc.add_paragraph()
    r  = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)


def _remove_table_borders(table):
    """Quita el estilo de borde por defecto de una tabla."""
    tbl   = table._tbl
    tblPr = tbl.tblPr
    for tag in ('w:tblStyle', 'w:tblBorders'):
        el = tblPr.find(qn(tag))
        if el is not None:
            tblPr.remove(el)


def _campo_pagina(run):
    """Inserta campo PAGE de Word en un run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instr)
    run._r.append(fldChar2)
    
def _sin_acentos(texto: str) -> str:
    """Normaliza para comparación case/accent-insensitive, sin alterar el texto visible."""
    return ''.join(
        c for c in unicodedata.normalize('NFD', texto)
        if unicodedata.category(c) != 'Mn'
    ).upper()

# ══════════════════════════════════════════════════════════════════
# REPORT SERVICE
# ══════════════════════════════════════════════════════════════════

class ReportService:

    # ─────────────────────────────────────────────────────────────
    # Helpers de color
    # ─────────────────────────────────────────────────────────────

    @staticmethod
    def _texto_seguro(valor, default=''):
        """Normaliza valores None provenientes de columnas NULL en la DB."""
        if valor is None:
            return default
        return str(valor)
    
    @staticmethod
    def _agrupar_findings_por_check(findings):
        grupos = {}
        for f in findings:
            key = (f.get('check_id'), f.get('servicio'))
            if key not in grupos:
                grupos[key] = {
                    'titulo': f.get('titulo'),
                    'servicio': f.get('servicio'),
                    'severidad': f.get('severidad'),
                    'descripcion': f.get('descripcion'),
                    'condicion_logica': f.get('condicion_logica'),
                    'remediacion': f.get('remediacion'),
                    'referencia': f.get('referencia'),
                    'recursos': []
                }
            grupos[key]['recursos'].append({
                'resource_id': f.get('resource_id'),
                'estado': f.get('estado', 'ABIERTO'),
                'comment': f.get('finding_comment') or '',
                'inventory_data': f.get('inventory_data') or '',   # ← nuevo
                'evidencias': f.get('evidencias', [])
            })
        return list(grupos.values())

    @staticmethod
    def _normalizar_color(color: str) -> str | None:
        if not color:
            return None
        color = color.strip().lstrip('#')
        if len(color) == 3:
            color = ''.join(c * 2 for c in color)
        if len(color) == 6:
            return f"FF{color.upper()}"
        if len(color) == 8:
            return color.upper()
        return None

    @staticmethod
    def construir_mapa_severidad(severidades):
        return {row["nombre"].upper(): row["color"] for row in severidades}

    @staticmethod
    def _limpiar(texto):
        if not texto:
            return "sin_valor"
        return re.sub(r'[^a-zA-Z0-9_-]', '_', str(texto))

    @staticmethod
    def generar_nombre_archivo(data, proyecto_id, extension="xlsx"):
        if data:
            titulo      = ReportService._limpiar(data[0].get('proyecto_titulo'))
            proveedores = {row.get('proveedor', 'sin_proveedor') for row in data}
            proveedor   = ReportService._limpiar("_".join(proveedores))
        else:
            titulo    = f"proyecto_{proyecto_id}"
            proveedor = "sin_proveedor"
        return f"{titulo}_{proveedor}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{extension}"

    # ─────────────────────────────────────────────────────────────
    # XLSX
    # ─────────────────────────────────────────────────────────────

    @staticmethod
    def generar_xlsx(data, severidades):
        wb = Workbook()
        ws = wb.active
        ws.title = "Findings"

        headers = [
            'proveedor', 'servicio', 'check_id', 'titulo', 'descripcion',
            'riesgo', 'condicion logica', 'remediacion', 'referencia',
            'resource_id', 'estado'
        ]
        ws.append(headers)

        wrap        = Alignment(wrap_text=True, vertical="top")
        header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)
        sev_colors  = ReportService.construir_mapa_severidad(severidades)

        for cell in ws[1]:
            cell.fill      = header_fill
            cell.font      = header_font
            cell.alignment = wrap

        for row in data:
            fila = [
                row.get('proveedor', ''),      row.get('servicio', ''),
                row.get('check_id', ''),       row.get('titulo', ''),
                row.get('descripcion', ''),    row.get('severidad', ''),
                row.get('condicion_logica', ''), row.get('remediacion', ''),
                row.get('referencia', ''),     row.get('resource_id', ''),
                row.get('estado', '')
            ]
            ws.append(fila)
            current_row = ws.max_row

            for col_idx, value in enumerate(fila, start=1):
                cell           = ws.cell(row=current_row, column=col_idx)
                cell.alignment = wrap
                if col_idx == 6:
                    raw   = sev_colors.get(str(value).upper())
                    color = ReportService._normalizar_color(raw)
                    if color:
                        cell.fill = PatternFill(start_color=color, end_color=color, fill_type="solid")

        ws.auto_filter.ref = ws.dimensions
        ws.freeze_panes    = "A2"

        for col in ws.columns:
            max_len    = max((len(str(c.value)) for c in col if c.value), default=0)
            col_letter = col[0].column_letter
            ws.column_dimensions[col_letter].width = min(max_len + 2, 50)

        output = BytesIO()
        wb.save(output)
        output.seek(0)
        return output

    # ─────────────────────────────────────────────────────────────
    # CSV VULMA
    # ─────────────────────────────────────────────────────────────

    @staticmethod
    def generar_csv_vulma(data):
        output = io.StringIO(newline='')
        writer = csv.DictWriter(
            output,
            fieldnames=[
                "Vulnerabilidad", "Descripcion", "Impacto", "Solucion",
                "Target", "Severidad", "CVSS", "CVSS_Vector", "Referencias",
                "CVE", "Exploits", "Output", "Protocolo", "Puerto", "URI"
            ],
            delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL
        )
        writer.writeheader()
        for row in data:
            writer.writerow({
                "Vulnerabilidad": row.get("titulo", ""),
                "Descripcion":    row.get("descripcion", ""),
                "Impacto":        row.get("condicion_logica", ""),
                "Solucion":       row.get("remediacion", ""),
                "Target":         row.get("resource_id", ""),
                "Severidad":      row.get("severidad", ""),
                "CVSS": "", "CVSS_Vector": "",
                "Referencias":    row.get("referencia", ""),
                "CVE": "", "Exploits": "", "Output": "",
                "Protocolo": "", "Puerto": "", "URI": ""
            })
        csv_content  = output.getvalue()
        output.close()
        final_output = io.BytesIO()
        final_output.write(csv_content.encode("utf-8-sig"))
        final_output.seek(0)
        return final_output

    # ─────────────────────────────────────────────────────────────
    # DOCX — bloques internos
    # ─────────────────────────────────────────────────────────────

    @staticmethod
    def _doc_base(proyecto, tema):
        doc = Document()

        # ── Forzar formato moderno (evita "Modo de compatibilidad") ──
        settings = doc.settings.element
        compat   = OxmlElement('w:compat')
        cs       = OxmlElement('w:compatSetting')
        cs.set(qn('w:name'), 'compatibilityMode')
        cs.set(qn('w:uri'),  'http://schemas.microsoft.com/office/word')
        cs.set(qn('w:val'),  '15')
        compat.append(cs)
        settings.append(compat)
        sec = doc.sections[0]
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

        # Fuente base
        doc.styles['Normal'].font.name       = 'Arial'
        doc.styles['Normal'].font.size       = Pt(10)
        doc.styles['Normal'].font.color.rgb  = _hex_to_rgb(tema.get('texto_oscuro', '#111827'))

        # ── Header ──
        header = sec.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.clear()

        r1 = hp.add_run(f"RedScope  |  {proyecto.get('titulo', '')}")
        r1.font.name      = 'Arial'
        r1.font.size      = Pt(8)
        r1.font.color.rgb = _hex_to_rgb(tema.get('borde', '#CCCCCC'))

        hp.add_run('\t')

        r2 = hp.add_run('CONFIDENCIAL')
        r2.font.name      = 'Arial'
        r2.font.size      = Pt(8)
        r2.font.bold      = True
        r2.font.color.rgb = _hex_to_rgb(tema.get('acento', '#00B4D8'))

        # Línea baja header
        pPr  = hp._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), tema.get('fondo_secundario', '#2D2A6E').lstrip('#'))
        pBdr.append(bot)
        pPr.append(pBdr)

        # ── Footer ──
        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.clear()

        rf = fp.add_run(f"{proyecto.get('cliente', '')}  ·  {datetime.now().strftime('%d/%m/%Y')}")
        rf.font.name      = 'Arial'
        rf.font.size      = Pt(8)
        rf.font.color.rgb = _hex_to_rgb(tema.get('borde', '#CCCCCC'))

        fp.add_run('\t')

        rp = fp.add_run('Página ')
        rp.font.name      = 'Arial'
        rp.font.size      = Pt(8)
        rp.font.color.rgb = _hex_to_rgb(tema.get('borde', '#CCCCCC'))

        rnum = fp.add_run()
        rnum.font.name      = 'Arial'
        rnum.font.size      = Pt(8)
        rnum.font.color.rgb = _hex_to_rgb(tema.get('borde', '#CCCCCC'))
        _campo_pagina(rnum)

        return doc

    @staticmethod
    def _bloque_portada(doc, proyecto, tema):
        """Portada: bloque de color + datos del proyecto."""
        color_primario    = tema.get('fondo_primario',   '#1E1B4B').lstrip('#')
        color_secundario  = tema.get('fondo_secundario', '#2D2A6E').lstrip('#')
        color_acento      = tema.get('acento',           '#00B4D8').lstrip('#')
        color_texto_claro = tema.get('texto_claro',      '#FFFFFF').lstrip('#')
        proveedor         = proyecto.get('tipo_servicio', 'Cloud').upper()
        cliente           = proyecto.get('cliente', '').upper()

        # ── Bloque hero ──
        t = doc.add_table(rows=1, cols=1)
        _remove_table_borders(t)
        c = t.cell(0, 0)
        _set_cell_bg(c, color_primario)
        _set_cell_margin(c, 600, 500, 400, 400)
        _set_row_height(t.rows[0], 8)

        # Nombre empresa pequeño arriba
        p0 = c.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run('P E R S O N A L')
        r0.font.name      = 'Arial'
        r0.font.size      = Pt(11)
        r0.font.bold      = True
        r0.font.color.rgb = _hex_to_rgb(color_acento)

        # H1 — Pentest Cloud — AWS
        p1 = c.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(18)
        r1 = p1.add_run(f'Pentest Cloud — {proveedor}')
        r1.font.name      = 'Arial'
        r1.font.size      = Pt(28)
        r1.font.bold      = True
        r1.font.color.rgb = _hex_to_rgb(color_texto_claro)

        # Subtítulo — nombre del cliente
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(10)
        r2 = p2.add_run(cliente)
        r2.font.name      = 'Arial'
        r2.font.size      = Pt(20)
        r2.font.bold      = True
        r2.font.color.rgb = _hex_to_rgb(color_acento)

        # ── Banda acento cyan ──
        banda = doc.add_table(rows=1, cols=1)
        _remove_table_borders(banda)
        bc = banda.cell(0, 0)
        _set_cell_bg(bc, color_acento)
        _set_row_height(banda.rows[0], 0.35)
        bc.paragraphs[0].add_run('')

        doc.add_paragraph()

        # ── Tabla de datos del proyecto ──
        campos = [
            ('Cliente',   proyecto.get('cliente',   '')),
            ('Proyecto',  proyecto.get('titulo',    '')),
            ('Proveedor', proveedor),
            ('Cuenta',    proyecto.get('cuenta_id', 'N/A')),
            ('Fecha',     datetime.now().strftime('%d de %B de %Y')),
        ]

        info = doc.add_table(rows=len(campos), cols=2)
        _remove_table_borders(info)
        info.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, (label, valor) in enumerate(campos):
            # Label
            lc = info.rows[i].cells[0]
            lc.width = Cm(4)
            _set_cell_bg(lc, color_secundario)
            _set_cell_borders(lc, 'FFFFFF', '2')
            _set_cell_margin(lc, 100, 100, 160, 120)
            lp = lc.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            lr = lp.add_run(label)
            lr.font.name      = 'Arial'
            lr.font.size      = Pt(10)
            lr.font.bold      = True
            lr.font.color.rgb = _hex_to_rgb(color_acento)
            lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Valor
            vc = info.rows[i].cells[1]
            vc.width = Cm(10)
            _set_cell_bg(vc, color_primario)
            _set_cell_borders(vc, 'FFFFFF', '2')
            _set_cell_margin(vc, 100, 100, 160, 120)
            vp = vc.paragraphs[0]
            vr = vp.add_run(str(valor))
            vr.font.name      = 'Arial'
            vr.font.size      = Pt(10)
            vr.font.color.rgb = _hex_to_rgb(color_texto_claro)
            vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        _page_break(doc)

    @staticmethod
    def _bloque_toc(doc, estructura, tema):
        """
        TOC nativo de Word con campo TOC 1.
        - El analista abre el doc y hace clic en 'Actualizar tabla'
        - Los números de página se actualizan automáticamente
        - Requiere que _seccion_titulo use estilo Heading para que Word lo detecte
        """
        color_primario = tema.get('fondo_primario', '#1E1B4B')
        color_acento   = tema.get('acento',         '#00B4D8')

        # ── Título ───────────────────────────────────────────────
        p_titulo = doc.add_paragraph()
        p_titulo.paragraph_format.space_before = Pt(0)
        p_titulo.paragraph_format.space_after  = Pt(4)
        r = p_titulo.add_run('Tabla de Contenidos')
        r.font.name      = 'Arial'
        r.font.size      = Pt(18)
        r.font.bold      = True
        r.font.color.rgb = _hex_to_rgb(color_primario)

        _add_hr(doc, color_acento.lstrip('#'), 12)

        # ── Campo TOC nativo de Word ──────────────────────────────
        # Word detecta párrafos con estilo Heading 1 y construye el índice
        p_toc = doc.add_paragraph()

        # Instrucción del campo: TOC con nivel 1, sin hiperlinks, con líderes
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        fldChar_begin.set(qn('w:dirty'), 'true')   # fuerza re-renderizado al abrir

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-1" \\h \\z \\u '
        # \\o "1-1" → solo Heading 1
        # \\h       → hipervínculos (Ctrl+clic navega a la sección)
        # \\z       → oculta números de página en web layout
        # \\u       → usa estilos de párrafo con outlineLevel

        fldChar_sep = OxmlElement('w:fldChar')
        fldChar_sep.set(qn('w:fldCharType'), 'separate')

        # Texto placeholder que ve el usuario antes de actualizar
        r_placeholder = OxmlElement('w:r')
        t_placeholder = OxmlElement('w:t')
        t_placeholder.text = '[Haga clic en Actualizar tabla para generar el índice]'
        r_placeholder.append(t_placeholder)

        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')

        # Ensamblar el campo en el run del párrafo
        run = p_toc.add_run()
        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_sep)
        run._r.append(r_placeholder)
        run._r.append(fldChar_end)

        _page_break(doc)

    @staticmethod
    def _seccion_titulo(doc, texto, tema):
        color_primario = tema.get('fondo_primario', '#1E1B4B')
        color_acento   = tema.get('acento',         '#00B4D8')

        p = doc.add_paragraph(style='Heading 1')  # ← clave para el TOC
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)

        r = p.add_run(texto)
        r.font.name      = 'Arial'
        r.font.size      = Pt(15)
        r.font.bold      = True
        r.font.color.rgb = _hex_to_rgb(color_primario)

        _add_hr(doc, color_acento.lstrip('#'), 10)
        
    @staticmethod
    def _bloque_detalle_vulnerabilidades(doc, findings, tema, severidades):
        """Una ficha por cada vulnerabilidad única (sin listar recursos)."""
        ReportService._seccion_titulo(doc, 'Detalle de Vulnerabilidades', tema)

        color_primario    = tema.get('fondo_primario',       '#1E1B4B')
        color_secundario  = tema.get('fondo_secundario',     '#2D2A6E')
        color_acento      = tema.get('acento',               '#00B4D8')
        color_texto_claro = tema.get('texto_claro',          '#FFFFFF')
        color_fila_par    = tema.get('fondo_tabla_fila_par', '#E8EAF6')
        color_oscuro      = tema.get('texto_oscuro',         '#111827')
        sev_map           = {s['nombre'].upper(): s['color'].lstrip('#') for s in severidades}

        intro = doc.add_paragraph()
        intro.paragraph_format.space_after = Pt(12)
        ri = intro.add_run(
            'A continuación se describe cada vulnerabilidad identificada durante la evaluación, '
            'incluyendo descripción técnica, condición lógica de detección, remediación recomendada '
            'y referencias asociadas. El detalle de los recursos afectados por cada una se presenta '
            'en la sección siguiente.'
        )
        ri.font.name      = 'Arial'
        ri.font.size      = Pt(10)
        ri.font.color.rgb = _hex_to_rgb(color_oscuro)

        grupos = ReportService._agrupar_findings_por_check(findings)

        for idx, g in enumerate(grupos, start=1):
            sev     = ReportService._texto_seguro(g.get('severidad')).upper()
            sev_hex = sev_map.get(sev, 'CCCCCC')

            enc = doc.add_table(rows=1, cols=2)
            _remove_table_borders(enc)

            tc_titulo = enc.rows[0].cells[0]
            tc_titulo.width = Cm(13)
            _set_cell_bg(tc_titulo, color_primario.lstrip('#'))
            _set_cell_borders(tc_titulo, 'FFFFFF', '2')
            _set_cell_margin(tc_titulo, 120, 120, 160, 160)
            tp = tc_titulo.paragraphs[0]
            tp.paragraph_format.space_before = Pt(4)
            rn = tp.add_run(f'#{idx}  ')
            rn.font.name = 'Arial'; rn.font.size = Pt(11); rn.font.bold = True
            rn.font.color.rgb = _hex_to_rgb(color_acento)
            rt = tp.add_run(ReportService._texto_seguro(g.get('titulo')))
            rt.font.name = 'Arial'; rt.font.size = Pt(11); rt.font.bold = True
            rt.font.color.rgb = _hex_to_rgb(color_texto_claro)
            tc_titulo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            tc_sev = enc.rows[0].cells[1]
            tc_sev.width = Cm(4)
            _set_cell_bg(tc_sev, sev_hex)
            _set_cell_borders(tc_sev, 'FFFFFF', '2')
            _set_cell_margin(tc_sev)
            sp = tc_sev.paragraphs[0]
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sp.add_run(sev)
            sr.font.name = 'Arial'; sr.font.size = Pt(11); sr.font.bold = True
            sr.font.color.rgb = _hex_to_rgb(color_texto_claro)
            tc_sev.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            meta = doc.add_table(rows=1, cols=1)
            _remove_table_borders(meta)
            mc = meta.rows[0].cells[0]
            mc.width = Cm(17)
            _set_cell_bg(mc, color_fila_par.lstrip('#'))
            _set_cell_borders(mc, 'CCCCCC', '2')
            _set_cell_margin(mc)
            mp = mc.paragraphs[0]
            mr_label = mp.add_run('Servicio: ')
            mr_label.font.name = 'Arial'; mr_label.font.size = Pt(9); mr_label.font.bold = True
            mr_label.font.color.rgb = _hex_to_rgb(color_secundario)
            mr_val = mp.add_run(ReportService._texto_seguro(g.get('servicio')))
            mr_val.font.name = 'Arial'; mr_val.font.size = Pt(9)
            mr_val.font.color.rgb = _hex_to_rgb(color_oscuro)
            mc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            ficha_campos = [
                ('Descripción', ReportService._texto_seguro(g.get('descripcion'))),
                ('Condición',   ReportService._texto_seguro(g.get('condicion_logica'))),
                ('Remediación', ReportService._texto_seguro(g.get('remediacion'))),
                ('Referencia',  ReportService._texto_seguro(g.get('referencia'))),
            ]
            ficha = doc.add_table(rows=len(ficha_campos), cols=2)
            _remove_table_borders(ficha)
            for fi, (fl, fv) in enumerate(ficha_campos):
                bg = color_fila_par.lstrip('#') if fi % 2 == 0 else 'FFFFFF'
                lc = ficha.rows[fi].cells[0]
                lc.width = Cm(3.5)
                _set_cell_bg(lc, color_secundario.lstrip('#'))
                _set_cell_borders(lc, 'FFFFFF', '2')
                _set_cell_margin(lc)
                lp = lc.paragraphs[0]
                lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                lr = lp.add_run(fl)
                lr.font.name = 'Arial'; lr.font.size = Pt(9); lr.font.bold = True
                lr.font.color.rgb = _hex_to_rgb(color_acento)
                lc.vertical_alignment = WD_ALIGN_VERTICAL.TOP

                vc = ficha.rows[fi].cells[1]
                vc.width = Cm(13.5)
                _set_cell_bg(vc, bg)
                _set_cell_borders(vc, 'DDDDDD', '2')
                _set_cell_margin(vc)
                vp = vc.paragraphs[0]
                vr = vp.add_run(fv)
                vr.font.name = 'Arial'; vr.font.size = Pt(9)
                vr.font.color.rgb = _hex_to_rgb(color_oscuro)
                vc.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            doc.add_paragraph()
            if idx < len(grupos):
                _page_break(doc)

    @staticmethod
    def _bloque_seccion_estatica(doc, subtitulo, contenido, tema):
        ReportService._seccion_titulo(doc, subtitulo, tema)

        texto = contenido or '[Completar por el analista]'
        texto = texto.replace('\r\n', '\n').replace('\r', '\n')  # normaliza CRLF/CR sueltos
        parrafos = [p.strip() for p in texto.split('\n\n') if p.strip()]

        color = tema.get('texto_oscuro', '#111827') if contenido else tema.get('borde', '#CCCCCC')

        for parrafo_texto in parrafos:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            lineas = parrafo_texto.split('\n')
            for i, linea in enumerate(lineas):
                r = p.add_run(linea)
                r.font.name      = 'Arial'
                r.font.size      = Pt(10)
                r.font.italic    = contenido is None
                r.font.color.rgb = _hex_to_rgb(color)
                if i < len(lineas) - 1:
                    r.add_break()

        doc.add_paragraph()
        
    @staticmethod
    def _bloque_seccion_vacia(doc, subtitulo, tema):
        """Sección estática — el analista la completa a mano."""
        ReportService._seccion_titulo(doc, subtitulo, tema)
        p = doc.add_paragraph()
        r = p.add_run('[Completar por el analista]')
        r.font.name      = 'Arial'
        r.font.size      = Pt(10)
        r.font.italic    = True
        r.font.color.rgb = _hex_to_rgb(tema.get('borde', '#CCCCCC'))
        doc.add_paragraph()

    @staticmethod
    def _bloque_resumen(doc, findings, tema, severidades):
        """Tabla resumen de vulnerabilidades únicas por severidad."""
        ReportService._seccion_titulo(doc, 'Resumen de Hallazgos', tema)

        color_primario    = tema.get('fondo_primario',      '#1E1B4B')
        color_texto_claro = tema.get('texto_claro',         '#FFFFFF')
        color_fila_par    = tema.get('fondo_tabla_fila_par','#E8EAF6')
        color_oscuro      = tema.get('texto_oscuro',        '#111827')

        grupos = ReportService._agrupar_findings_por_check(findings)

        sev_map = {s['nombre'].upper(): s['color'].lstrip('#') for s in severidades}
        conteo = {s['nombre'].upper(): 0 for s in sorted(severidades, key=lambda x: x['orden'], reverse=True)}
        sin_clasificar = 0

        for g in grupos:
            sev = ReportService._texto_seguro(g.get('severidad')).upper()
            if sev in conteo:
                conteo[sev] += 1
            else:
                sin_clasificar += 1

        total_vulnerabilidades = len(grupos)
        total_recursos = sum(len(g['recursos']) for g in grupos)

        intro = doc.add_paragraph()
        ri = intro.add_run(
            f"Durante la evaluación se identificaron {total_vulnerabilidades} vulnerabilidad(es) "
            f"distintas, afectando un total de {total_recursos} recurso(s), distribuidas en las "
            f"siguientes categorías de riesgo:"
        )
        ri.font.name      = 'Arial'
        ri.font.size      = Pt(10)
        ri.font.color.rgb = _hex_to_rgb(color_oscuro)
        intro.paragraph_format.space_after = Pt(8)

        table = doc.add_table(rows=1, cols=3)
        _remove_table_borders(table)

        for ci, txt in enumerate(['Severidad', 'Cantidad', 'Porcentaje']):
            c = table.rows[0].cells[ci]
            _set_cell_bg(c, color_primario.lstrip('#'))
            _set_cell_borders(c, 'FFFFFF', '2')
            _set_cell_margin(c)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(txt)
            r.font.name      = 'Arial'
            r.font.size      = Pt(10)
            r.font.bold      = True
            r.font.color.rgb = _hex_to_rgb(color_texto_claro)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for sev, qty in conteo.items():
            if qty == 0:
                continue
            pct      = f"{(qty / total_vulnerabilidades * 100):.0f}%" if total_vulnerabilidades > 0 else "0%"
            sev_hex  = sev_map.get(sev, 'CCCCCC')
            row      = table.add_row()

            sc = row.cells[0]
            _set_cell_bg(sc, sev_hex)
            _set_cell_borders(sc, 'FFFFFF', '2')
            _set_cell_margin(sc)
            sp = sc.paragraphs[0]
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sp.add_run(sev)
            sr.font.name      = 'Arial'
            sr.font.size      = Pt(10)
            sr.font.bold      = True
            sr.font.color.rgb = _hex_to_rgb(color_texto_claro)
            sc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            qc = row.cells[1]
            _set_cell_bg(qc, color_fila_par.lstrip('#'))
            _set_cell_borders(qc, 'CCCCCC', '2')
            _set_cell_margin(qc)
            qp = qc.paragraphs[0]
            qp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            qr = qp.add_run(str(qty))
            qr.font.name      = 'Arial'
            qr.font.size      = Pt(10)
            qr.font.bold      = True
            qr.font.color.rgb = _hex_to_rgb(color_oscuro)
            qc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            pc = row.cells[2]
            _set_cell_bg(pc, color_fila_par.lstrip('#'))
            _set_cell_borders(pc, 'CCCCCC', '2')
            _set_cell_margin(pc)
            pp = pc.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pr = pp.add_run(pct)
            pr.font.name      = 'Arial'
            pr.font.size      = Pt(10)
            pr.font.color.rgb = _hex_to_rgb(color_oscuro)
            pc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        if sin_clasificar > 0:
            pct = f"{(sin_clasificar / total_vulnerabilidades * 100):.0f}%" if total_vulnerabilidades > 0 else "0%"
            row = table.add_row()

            sc = row.cells[0]
            _set_cell_bg(sc, 'CCCCCC')
            _set_cell_borders(sc, 'FFFFFF', '2')
            _set_cell_margin(sc)
            sp = sc.paragraphs[0]
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sp.add_run('SIN CLASIFICAR')
            sr.font.name      = 'Arial'
            sr.font.size      = Pt(10)
            sr.font.bold      = True
            sr.font.color.rgb = _hex_to_rgb(color_oscuro)
            sc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            qc = row.cells[1]
            _set_cell_bg(qc, color_fila_par.lstrip('#'))
            _set_cell_borders(qc, 'CCCCCC', '2')
            _set_cell_margin(qc)
            qp = qc.paragraphs[0]
            qp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            qr = qp.add_run(str(sin_clasificar))
            qr.font.name      = 'Arial'
            qr.font.size      = Pt(10)
            qr.font.bold      = True
            qr.font.color.rgb = _hex_to_rgb(color_oscuro)
            qc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            pc = row.cells[2]
            _set_cell_bg(pc, color_fila_par.lstrip('#'))
            _set_cell_borders(pc, 'CCCCCC', '2')
            _set_cell_margin(pc)
            pp = pc.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pr = pp.add_run(pct)
            pr.font.name      = 'Arial'
            pr.font.size      = Pt(10)
            pr.font.color.rgb = _hex_to_rgb(color_oscuro)
            pc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        doc.add_paragraph()

    @staticmethod
    def _bloque_tabla_hallazgos(doc, findings, tema, severidades):
        """Tabla índice de todos los hallazgos."""
        ReportService._seccion_titulo(doc, 'Hallazgos', tema)

        color_primario    = tema.get('fondo_primario',       '#1E1B4B')
        color_texto_claro = tema.get('texto_claro',          '#FFFFFF')
        color_fila_par    = tema.get('fondo_tabla_fila_par', '#E8EAF6')
        color_oscuro      = tema.get('texto_oscuro',         '#111827')
        sev_map           = {s['nombre'].upper(): s['color'].lstrip('#') for s in severidades}

        cols    = ['#', 'Título', 'Servicio', 'Recurso', 'Severidad']
        widths  = [1.0, 5.5, 2.5, 5.0, 2.5]

        table = doc.add_table(rows=1, cols=len(cols))
        _remove_table_borders(table)

        for ci, (txt, w) in enumerate(zip(cols, widths)):
            c = table.rows[0].cells[ci]
            c.width = Cm(w)
            _set_cell_bg(c, color_primario.lstrip('#'))
            _set_cell_borders(c, 'FFFFFF', '2')
            _set_cell_margin(c)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(txt)
            r.font.name      = 'Arial'
            r.font.size      = Pt(9)
            r.font.bold      = True
            r.font.color.rgb = _hex_to_rgb(color_texto_claro)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for idx, f in enumerate(findings, start=1):
            sev     = ReportService._texto_seguro(f.get('severidad')).upper()
            sev_hex = sev_map.get(sev, 'CCCCCC')
            bg      = color_fila_par.lstrip('#') if idx % 2 == 0 else 'FFFFFF'
            row     = table.add_row()

            valores = [
                str(idx),
                ReportService._texto_seguro(f.get('titulo')),
                ReportService._texto_seguro(f.get('servicio')),
                ReportService._texto_seguro(f.get('resource_id')),
                sev,
            ]

            for ci, (val, w) in enumerate(zip(valores, widths)):
                c = row.cells[ci]
                c.width = Cm(w)
                _set_cell_margin(c)
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                if ci == 4:
                    _set_cell_bg(c, sev_hex)
                    _set_cell_borders(c, 'FFFFFF', '2')
                    p = c.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(val)
                    r.font.name      = 'Arial'
                    r.font.size      = Pt(9)
                    r.font.bold      = True
                    r.font.color.rgb = _hex_to_rgb(color_texto_claro)
                else:
                    _set_cell_bg(c, bg)
                    _set_cell_borders(c, 'DDDDDD', '2')
                    p = c.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    r = p.add_run(val)
                    r.font.name      = 'Arial'
                    r.font.size      = Pt(9)
                    r.font.color.rgb = _hex_to_rgb(color_oscuro)

        doc.add_paragraph()

    @staticmethod
    def _bloque_detalle_hallazgos(doc, findings, tema, severidades, base_dir=None):
        """Una ficha completa por cada vulnerabilidad única, con sus recursos afectados y evidencia."""
        ReportService._seccion_titulo(doc, 'Detalle de Hallazgos', tema)
        color_primario    = tema.get('fondo_primario',       '#1E1B4B')
        color_secundario  = tema.get('fondo_secundario',     '#2D2A6E')
        color_acento      = tema.get('acento',               '#00B4D8')
        color_texto_claro = tema.get('texto_claro',          '#FFFFFF')
        color_fila_par    = tema.get('fondo_tabla_fila_par', '#E8EAF6')
        color_oscuro      = tema.get('texto_oscuro',         '#111827')
        sev_map           = {s['nombre'].upper(): s['color'].lstrip('#') for s in severidades}

        intro = doc.add_paragraph()
        intro.paragraph_format.space_after = Pt(12)
        ri = intro.add_run(
            'A continuación se presenta el detalle técnico de cada vulnerabilidad identificada '
            'durante la evaluación, incluyendo descripción, condición lógica, remediación, '
            'referencias asociadas, recursos afectados y evidencia de las pruebas realizadas.'
        )
        ri.font.name      = 'Arial'
        ri.font.size      = Pt(10)
        ri.font.color.rgb = _hex_to_rgb(color_oscuro)

        grupos = ReportService._agrupar_findings_por_check(findings)

        for idx, g in enumerate(grupos, start=1):
            sev      = ReportService._texto_seguro(g.get('severidad')).upper()
            sev_hex  = sev_map.get(sev, 'CCCCCC')
            recursos = g['recursos']

            # ── Servicio (línea suelta, arriba de todo) ───────────
            p_servicio = doc.add_paragraph()
            p_servicio.paragraph_format.space_after = Pt(2)
            rs_label = p_servicio.add_run('Servicio: ')
            rs_label.font.name = 'Arial'; rs_label.font.size = Pt(9); rs_label.font.bold = True
            rs_label.font.color.rgb = _hex_to_rgb(color_secundario)
            rs_val = p_servicio.add_run(ReportService._texto_seguro(g.get('servicio')).upper())
            rs_val.font.name = 'Arial'; rs_val.font.size = Pt(10); rs_val.font.bold = True
            rs_val.font.color.rgb = _hex_to_rgb(color_primario)

            # ── Encabezado: título + severidad ────────────────────
            enc = doc.add_table(rows=1, cols=2)
            _remove_table_borders(enc)

            tc_titulo = enc.rows[0].cells[0]
            tc_titulo.width = Cm(13)
            _set_cell_bg(tc_titulo, color_primario.lstrip('#'))
            _set_cell_borders(tc_titulo, 'FFFFFF', '2')
            _set_cell_margin(tc_titulo, 120, 120, 160, 160)
            tp = tc_titulo.paragraphs[0]
            tp.paragraph_format.space_before = Pt(4)
            rn = tp.add_run(f'#{idx}  ')
            rn.font.name = 'Arial'; rn.font.size = Pt(11); rn.font.bold = True
            rn.font.color.rgb = _hex_to_rgb(color_acento)
            rt = tp.add_run(ReportService._texto_seguro(g.get('titulo')))
            rt.font.name = 'Arial'; rt.font.size = Pt(11); rt.font.bold = True
            rt.font.color.rgb = _hex_to_rgb(color_texto_claro)
            tc_titulo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            tc_sev = enc.rows[0].cells[1]
            tc_sev.width = Cm(4)
            _set_cell_bg(tc_sev, sev_hex)
            _set_cell_borders(tc_sev, 'FFFFFF', '2')
            _set_cell_margin(tc_sev)
            sp = tc_sev.paragraphs[0]
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sp.add_run(sev)
            sr.font.name = 'Arial'; sr.font.size = Pt(11); sr.font.bold = True
            sr.font.color.rgb = _hex_to_rgb(color_texto_claro)
            tc_sev.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # ── Campos técnicos ────────────────────────────────────
            ficha_campos = [
                ('Descripción', ReportService._texto_seguro(g.get('descripcion'))),
                ('Condición',   ReportService._texto_seguro(g.get('condicion_logica'))),
                ('Remediación', ReportService._texto_seguro(g.get('remediacion'))),
                ('Referencia',  ReportService._texto_seguro(g.get('referencia'))),
            ]
            ficha = doc.add_table(rows=len(ficha_campos), cols=2)
            _remove_table_borders(ficha)
            for fi, (fl, fv) in enumerate(ficha_campos):
                bg = color_fila_par.lstrip('#') if fi % 2 == 0 else 'FFFFFF'
                lc = ficha.rows[fi].cells[0]
                lc.width = Cm(3.5)
                _set_cell_bg(lc, color_secundario.lstrip('#'))
                _set_cell_borders(lc, 'FFFFFF', '2')
                _set_cell_margin(lc)
                lp = lc.paragraphs[0]
                lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                lr = lp.add_run(fl)
                lr.font.name = 'Arial'; lr.font.size = Pt(9); lr.font.bold = True
                lr.font.color.rgb = _hex_to_rgb(color_acento)
                lc.vertical_alignment = WD_ALIGN_VERTICAL.TOP

                vc = ficha.rows[fi].cells[1]
                vc.width = Cm(13.5)
                _set_cell_bg(vc, bg)
                _set_cell_borders(vc, 'DDDDDD', '2')
                _set_cell_margin(vc)
                vp = vc.paragraphs[0]
                vr = vp.add_run(fv)
                vr.font.name = 'Arial'; vr.font.size = Pt(9)
                vr.font.color.rgb = _hex_to_rgb(color_oscuro)
                vc.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            # ── Tabla de recursos afectados ───────────────────────
            p_recursos = doc.add_paragraph()
            p_recursos.paragraph_format.space_before = Pt(10)
            p_recursos.paragraph_format.space_after  = Pt(4)
            rrec = p_recursos.add_run(f'Recursos Afectados: {len(recursos)}')
            rrec.font.name = 'Arial'; rrec.font.size = Pt(9); rrec.font.bold = True
            rrec.font.color.rgb = _hex_to_rgb(color_secundario)

            tabla_rec = doc.add_table(rows=1, cols=2)
            _remove_table_borders(tabla_rec)
            for ci, (txt, w) in enumerate([('Recurso', 13.5), ('Estado', 3.5)]):
                c = tabla_rec.rows[0].cells[ci]
                c.width = Cm(w)
                _set_cell_bg(c, color_primario.lstrip('#'))
                _set_cell_borders(c, 'FFFFFF', '2')
                _set_cell_margin(c)
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(txt)
                r.font.name = 'Arial'; r.font.size = Pt(9); r.font.bold = True
                r.font.color.rgb = _hex_to_rgb(color_texto_claro)
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            for ri_idx, rec in enumerate(recursos):
                bg = color_fila_par.lstrip('#') if ri_idx % 2 == 0 else 'FFFFFF'
                row = tabla_rec.add_row()

                c0 = row.cells[0]
                c0.width = Cm(13.5)
                _set_cell_bg(c0, bg)
                _set_cell_borders(c0, 'DDDDDD', '2')
                _set_cell_margin(c0)
                r0 = c0.paragraphs[0].add_run(ReportService._texto_seguro(rec['resource_id']))
                r0.font.name = 'Arial'; r0.font.size = Pt(9)
                r0.font.color.rgb = _hex_to_rgb(color_oscuro)

                c1 = row.cells[1]
                c1.width = Cm(3.5)
                _set_cell_bg(c1, bg)
                _set_cell_borders(c1, 'DDDDDD', '2')
                _set_cell_margin(c1)
                p1 = c1.paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r1 = p1.add_run(ReportService._texto_seguro(rec['estado'], 'ABIERTO'))
                r1.font.name = 'Arial'; r1.font.size = Pt(9)
                r1.font.color.rgb = _hex_to_rgb(color_oscuro)

            # ── Evidencia por recurso (solo los que tienen) ───────
            recursos_con_evidencia = [r for r in recursos if r['comment'] or r['evidencias'] or r.get('inventory_data')]
            if recursos_con_evidencia:
                p_ev_titulo = doc.add_paragraph()
                p_ev_titulo.paragraph_format.space_before = Pt(10)
                p_ev_titulo.paragraph_format.space_after  = Pt(4)
                rev_titulo = p_ev_titulo.add_run('Evidencia:')
                rev_titulo.font.name = 'Arial'; rev_titulo.font.size = Pt(9); rev_titulo.font.bold = True
                rev_titulo.font.color.rgb = _hex_to_rgb(color_secundario)

                for rec_idx, rec in enumerate(recursos_con_evidencia):
                    if rec_idx > 0:
                        _page_break(doc)

                    p_res = doc.add_paragraph()
                    p_res.paragraph_format.space_before = Pt(4)
                    r_res = p_res.add_run(f"• {rec['resource_id']}")
                    r_res.font.name = 'Arial'; r_res.font.size = Pt(9); r_res.font.bold = True
                    r_res.font.color.rgb = _hex_to_rgb(color_acento)

                    # ── Salida de la herramienta (solo si hay dato) ───────────
                    if rec.get('inventory_data'):
                        p_tool_label = doc.add_paragraph()
                        p_tool_label.paragraph_format.space_before = Pt(6)
                        p_tool_label.paragraph_format.space_after  = Pt(2)
                        rt_label = p_tool_label.add_run('Salida de la herramienta:')
                        rt_label.font.name = 'Arial'; rt_label.font.size = Pt(8); rt_label.font.bold = True
                        rt_label.font.color.rgb = _hex_to_rgb(color_secundario)

                        tabla_tool = doc.add_table(rows=1, cols=1)
                        _remove_table_borders(tabla_tool)
                        c_tool = tabla_tool.rows[0].cells[0]
                        _set_cell_bg(c_tool, '1E1E1E')
                        _set_cell_borders(c_tool, 'CCCCCC', '2')
                        _set_cell_margin(c_tool)
                        p_tool = c_tool.paragraphs[0]
                        r_tool = p_tool.add_run(rec['inventory_data'])
                        r_tool.font.name = 'Consolas'; r_tool.font.size = Pt(8)
                        r_tool.font.color.rgb = _hex_to_rgb('D4D4D4')

                    # ── Prueba manual: comentario + captura (solo si hay alguno) ──
                    if rec['comment'] or rec['evidencias']:
                        p_manual_label = doc.add_paragraph()
                        p_manual_label.paragraph_format.space_before = Pt(8)
                        p_manual_label.paragraph_format.space_after  = Pt(2)
                        rm_label = p_manual_label.add_run('Prueba Manual:')
                        rm_label.font.name = 'Arial'; rm_label.font.size = Pt(8); rm_label.font.bold = True
                        rm_label.font.color.rgb = _hex_to_rgb(color_secundario)

                        if rec['comment']:
                            p_comment = doc.add_paragraph()
                            p_comment.paragraph_format.space_after = Pt(4)
                            rc = p_comment.add_run(rec['comment'])
                            rc.font.name = 'Arial'; rc.font.size = Pt(9)
                            rc.font.color.rgb = _hex_to_rgb(color_oscuro)

                        if rec['evidencias'] and base_dir:
                            for img_path in rec['evidencias']:
                                abs_path = os.path.join(base_dir, img_path)
                                if os.path.exists(abs_path):
                                    try:
                                        p_img = doc.add_paragraph()
                                        p_img.paragraph_format.space_before = Pt(4)
                                        p_img.paragraph_format.space_after  = Pt(4)
                                        run_img = p_img.add_run()
                                        run_img.add_picture(abs_path, width=Cm(16.5))
                                    except Exception:
                                        pass

            doc.add_paragraph()
            if idx < len(grupos):
                _page_break(doc)

    # ─────────────────────────────────────────────────────────────
    # DOCX — punto de entrada público
    # ─────────────────────────────────────────────────────────────
    @staticmethod
    def generar_docx(data, proyecto, tema, estructura, severidades, contenido_secciones=None, base_dir=None):
        contenido_secciones = contenido_secciones or {}

        # Procesar imágenes que vienen como string separado por |
        for f in data:
            imagenes_raw  = f.get('imagenes') or ''
            f['evidencias'] = [img for img in imagenes_raw.split('|') if img]

        doc = ReportService._doc_base(proyecto, tema)
        ReportService._bloque_portada(doc, proyecto, tema)
        ReportService._bloque_toc(doc, estructura, tema)

        for seccion in estructura:
            tipo      = seccion['tipo']
            clave     = seccion['clave']
            subtitulo = seccion['subtitulo']
            dinamico  = seccion['es_dinamico']

            if tipo in ('portada', 'toc'):
                continue

            if dinamico:
                if clave == 'resumen_hallazgos':
                    ReportService._bloque_resumen(doc, data, tema, severidades)
                elif clave == 'hallazgos':
                    ReportService._bloque_tabla_hallazgos(doc, data, tema, severidades)
                elif clave == 'detalle_vulnerabilidades':
                    ReportService._bloque_detalle_vulnerabilidades(doc, data, tema, severidades)
                elif clave == 'detalle_hallazgos':
                    ReportService._bloque_detalle_hallazgos(doc, data, tema, severidades, base_dir=base_dir)
            else:
                contenido = contenido_secciones.get(clave)
                if clave == 'anexo_clasificacion':
                    ReportService._bloque_anexo_clasificacion(doc, contenido, tema, severidades)
                else:
                    ReportService._bloque_seccion_estatica(doc, subtitulo, contenido, tema)

            if seccion != estructura[-1]:
                _page_break(doc)

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output
        
    @staticmethod
    def _bloque_anexo_clasificacion(doc, contenido, tema, severidades):
        """Anexo de clasificación de riesgo con badges de color tomados dinámicamente de severidades."""
        ReportService._seccion_titulo(doc, 'Anexo 3: Clasificación del Riesgo', tema)

        color_oscuro      = tema.get('texto_oscuro', '#111827')
        color_texto_claro = tema.get('texto_claro',  '#FFFFFF')

        texto = (contenido or '').replace('\r\n', '\n').replace('\r', '\n')
        lineas = [l.strip() for l in texto.split('\n') if l.strip()]

        severidades_ordenadas = sorted(severidades, key=lambda s: s['orden'], reverse=True)
        intro_lineas = []
        descripciones = {}
        actual = None

        nombres = [s['nombre'].upper() for s in severidades_ordenadas]

        for linea in lineas:
            match = next((n for n in nombres if _sin_acentos(linea).startswith(_sin_acentos(n) + ':')), None)
            if match:
                actual = match
                descripciones[actual] = linea[len(match) + 1:].strip()
            elif actual:
                descripciones[actual] += ' ' + linea
            else:
                intro_lineas.append(linea)

        # Párrafo introductorio (todo lo que está antes de la primera severidad)
        if intro_lineas:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            r = p.add_run(' '.join(intro_lineas))
            r.font.name      = 'Arial'
            r.font.size      = Pt(10)
            r.font.color.rgb = _hex_to_rgb(color_oscuro)

        # Una fila badge + descripción por cada severidad activa en la DB
        for s in severidades_ordenadas:
            nombre      = s['nombre'].upper()
            color_hex   = s['color'].lstrip('#')
            descripcion = descripciones.get(nombre, '')

            fila = doc.add_table(rows=1, cols=2)
            _remove_table_borders(fila)

            c_badge = fila.rows[0].cells[0]
            c_badge.width = Cm(3.0)
            _set_cell_bg(c_badge, color_hex)
            _set_cell_borders(c_badge, 'FFFFFF', '2')
            _set_cell_margin(c_badge)
            pb = c_badge.paragraphs[0]
            pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rb = pb.add_run(nombre)
            rb.font.name      = 'Arial'
            rb.font.size      = Pt(9)
            rb.font.bold      = True
            rb.font.color.rgb = _hex_to_rgb(color_texto_claro)
            c_badge.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            c_desc = fila.rows[0].cells[1]
            c_desc.width = Cm(14.0)
            _set_cell_margin(c_desc)
            pd = c_desc.paragraphs[0]
            rd = pd.add_run(descripcion)
            rd.font.name      = 'Arial'
            rd.font.size      = Pt(9)
            rd.font.color.rgb = _hex_to_rgb(color_oscuro)
            c_desc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            sep = doc.add_paragraph()
            sep.paragraph_format.space_after = Pt(2)

        doc.add_paragraph()