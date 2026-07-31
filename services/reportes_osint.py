from io import BytesIO
from datetime import datetime
import re
import csv
import io
import unicodedata
import os

from openpyxl import Workbook
from openpyxl.styles import Alignment, PatternFill, Font

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Para reutilizar métodos dinámicos de reportes_aws
from services.reportes_aws import ReportService as ReportServiceAWS


# ══════════════════════════════════════════════════════════════════
# HELPERS XML (internos, no exponer fuera del módulo)
# ══════════════════════════════════════════════════════════════════
def _set_row_height(row, height_cm: float):
    """Fija la altura mínima de una fila en cm."""
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    for old in trPr.findall(qn('w:trHeight')):
        trPr.remove(old)
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'),   str(int(height_cm * 567)))
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)
def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Convierte '#1E1B4B' o '1E1B4B' a RGBColor."""
    h = hex_color.strip().lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_color: str):
    """Relleno sólido de celda."""
    h    = hex_color.strip().lstrip('#')
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # eliminar shd previo si existe
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  h.upper())
    tcPr.append(shd)


def _set_cell_borders(cell, hex_color: str = 'CCCCCC', size: str = '4'):
    """Bordes finos en todas las caras de la celda."""
    h    = hex_color.strip().lstrip('#').upper()
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    size)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), h)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_margin(cell, top=80, bottom=80, left=120, right=120):
    """Padding interno de celda en twips."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    mar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def _add_hr(doc, hex_color: str = '00B4D8', size: int = 10):
    """Línea horizontal decorativa debajo de un párrafo."""
    h = hex_color.strip().lstrip('#').upper()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(size))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), h)
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def _page_break(doc):
    p  = doc.add_paragraph()
    r  = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)


def _remove_table_borders(table):
    """Quita el estilo de borde por defecto de una tabla."""
    tbl   = table._tbl
    tblPr = tbl.tblPr
    for tag in ('w:tblStyle', 'w:tblBorders'):
        el = tblPr.find(qn(tag))
        if el is not None:
            tblPr.remove(el)


def _campo_pagina(run):
    """Inserta campo PAGE de Word en un run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instr)
    run._r.append(fldChar2)
    
def _sin_acentos(texto: str) -> str:
    """Normaliza para comparación case/accent-insensitive, sin alterar el texto visible."""
    return ''.join(
        c for c in unicodedata.normalize('NFD', texto)
        if unicodedata.category(c) != 'Mn'
    ).upper()


# ══════════════════════════════════════════════════════════════════
# REPORT SERVICE
# ══════════════════════════════════════════════════════════════════

class ReportService:

    TIPOS_INFORME = {
    'tecnico':   {'incluir_evidencia': True,  'excluir_secciones': [],'label': 'INFORME TÉCNICO'},
    'ejecutivo': {'incluir_evidencia': False, 'excluir_secciones': ['detalle_hallazgos'], 'label': 'INFORME EJECUTIVO'},
    }
    # ─────────────────────────────────────────────────────────────
    # Helpers de color
    # ─────────────────────────────────────────────────────────────
    @staticmethod
    def _texto_seguro(valor, default=''):
        """Normaliza valores None provenientes de columnas NULL en la DB."""
        if valor is None:
            return default
        return str(valor)
    

    @staticmethod
    def _normalizar_color(color: str) -> str | None:
        if not color:
            return None
        color = color.strip().lstrip('#')
        if len(color) == 3:
            color = ''.join(c * 2 for c in color)
        if len(color) == 6:
            return f"FF{color.upper()}"
        if len(color) == 8:
            return color.upper()
        return None

    @staticmethod
    def _limpiar(texto):
        if not texto:
            return "sin_valor"
        return re.sub(r'[^a-zA-Z0-9_-]', '_', str(texto))

    @staticmethod
    def generar_nombre_archivo(data, proyecto_id, extension="xlsx", tipo_informe=None, proyecto=None):
        cliente = ReportService._limpiar(proyecto.get('cliente', 'sin_cliente')) if proyecto else 'sin_cliente'
        tipo_proyecto = ReportService._limpiar(proyecto.get('tipo_proyecto', 'sin_tipo')) if proyecto else 'sin_tipo'
        tipo_servicio = ReportService._limpiar(proyecto.get('tipo_servicio', 'sin_servicio')) if proyecto else 'sin_servicio'

        sufijo_tipo = f"_{tipo_informe}" if tipo_informe else ""
        return f"{cliente}_{tipo_proyecto}_{tipo_servicio}{sufijo_tipo}.{extension}"

    # ─────────────────────────────────────────────────────────────
    # DOCX — bloques internos
    # ─────────────────────────────────────────────────────────────

    @staticmethod
    def _doc_base(proyecto, tema):
        doc = Document()

        # ── Forzar formato moderno (evita "Modo de compatibilidad") ──
        settings = doc.settings.element
        compat   = OxmlElement('w:compat')
        cs       = OxmlElement('w:compatSetting')
        cs.set(qn('w:name'), 'compatibilityMode')
        cs.set(qn('w:uri'),  'http://schemas.microsoft.com/office/word')
        cs.set(qn('w:val'),  '15')
        compat.append(cs)
        settings.append(compat)
        sec = doc.sections[0]
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

        # Fuente base
        doc.styles['Normal'].font.name       = 'Arial'
        doc.styles['Normal'].font.size       = Pt(10)
        doc.styles['Normal'].font.color.rgb  = _hex_to_rgb(tema.get('texto_oscuro', '#111827'))

        # ── Header ──
        header = sec.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.clear()

        r1 = hp.add_run(f"RedScope  |  {proyecto.get('titulo', '')}")
        r1.font.name      = 'Arial'
        r1.font.size      = Pt(8)
        r1.font.color.rgb = _hex_to_rgb(tema.get('borde', '#CCCCCC'))

        hp.add_run('\t')

        r2 = hp.add_run('CONFIDENCIAL')
        r2.font.name      = 'Arial'
        r2.font.size      = Pt(8)
        r2.font.bold      = True
        r2.font.color.rgb = _hex_to_rgb(tema.get('acento', '#00B4D8'))

        # Línea baja header
        pPr  = hp._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), tema.get('fondo_secundario', '#2D2A6E').lstrip('#'))
        pBdr.append(bot)
        pPr.append(pBdr)

        # ── Footer ──
        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.clear()

        rf = fp.add_run(f"{proyecto.get('cliente', '')}  ·  {datetime.now().strftime('%d/%m/%Y')}")
        rf.font.name      = 'Arial'
        rf.font.size      = Pt(8)
        rf.font.color.rgb = _hex_to_rgb(tema.get('borde', '#CCCCCC'))

        fp.add_run('\t')

        rp = fp.add_run('Página ')
        rp.font.name      = 'Arial'
        rp.font.size      = Pt(8)
        rp.font.color.rgb = _hex_to_rgb(tema.get('borde', '#CCCCCC'))

        rnum = fp.add_run()
        rnum.font.name      = 'Arial'
        rnum.font.size      = Pt(8)
        rnum.font.color.rgb = _hex_to_rgb(tema.get('borde', '#CCCCCC'))
        _campo_pagina(rnum)

        return doc

    @staticmethod
    def _bloque_portada(doc, proyecto, tema, tipo_informe='tecnico'):
        """Portada: bloque de color + datos del proyecto."""
        color_primario    = tema.get('fondo_primario',   '#1E1B4B').lstrip('#')
        color_secundario  = tema.get('fondo_secundario', '#2D2A6E').lstrip('#')
        color_acento      = tema.get('acento',           '#00B4D8').lstrip('#')
        color_texto_claro = tema.get('texto_claro',      '#FFFFFF').lstrip('#')
        tipo_proyecto     = proyecto.get('tipo_proyecto', 'OSINT').upper()
        tipo_servicio     = proyecto.get('tipo_servicio', 'N/A').upper()

        t = doc.add_table(rows=1, cols=1)
        _remove_table_borders(t)
        c = t.cell(0, 0)
        _set_cell_bg(c, color_primario)
        _set_cell_margin(c, 600, 500, 400, 400)
        _set_row_height(t.rows[0], 8)

        p0 = c.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run('P E R S O N A L')
        r0.font.name      = 'Arial'
        r0.font.size      = Pt(11)
        r0.font.bold      = True
        r0.font.color.rgb = _hex_to_rgb(color_acento)

        p1 = c.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(18)
        r1 = p1.add_run('OSINT')
        r1.font.name      = 'Arial'
        r1.font.size      = Pt(28)
        r1.font.bold      = True
        r1.font.color.rgb = _hex_to_rgb(color_texto_claro)

        # ── Tipo de informe ──
        label_informe = ReportService.TIPOS_INFORME.get(tipo_informe, {}).get('label', '')
        if label_informe:
            p25 = c.add_paragraph()
            p25.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p25.paragraph_format.space_before = Pt(12)
            r25 = p25.add_run(label_informe)
            r25.font.name      = 'Arial'
            r25.font.size      = Pt(18)
            r25.font.bold      = True
            r25.font.color.rgb = _hex_to_rgb(color_texto_claro)

        banda = doc.add_table(rows=1, cols=1)
        _remove_table_borders(banda)
        bc = banda.cell(0, 0)
        _set_cell_bg(bc, color_acento)
        _set_row_height(banda.rows[0], 0.35)
        bc.paragraphs[0].add_run('')

        doc.add_paragraph()

        campos = [
            ('Cliente',           proyecto.get('cliente',   '')),
            ('Proyecto',          proyecto.get('titulo',    '')),
            ('Tipo de Servicio',  tipo_servicio),
            ('Fecha',             datetime.now().strftime('%d de %B de %Y')),
        ]

        info = doc.add_table(rows=len(campos), cols=2)
        _remove_table_borders(info)
        info.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, (label, valor) in enumerate(campos):
            lc = info.rows[i].cells[0]
            lc.width = Cm(4)
            _set_cell_bg(lc, color_secundario)
            _set_cell_borders(lc, 'FFFFFF', '2')
            _set_cell_margin(lc, 100, 100, 160, 120)
            lp = lc.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            lr = lp.add_run(label)
            lr.font.name      = 'Arial'
            lr.font.size      = Pt(10)
            lr.font.bold      = True
            lr.font.color.rgb = _hex_to_rgb(color_acento)
            lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            vc = info.rows[i].cells[1]
            vc.width = Cm(10)
            _set_cell_bg(vc, color_primario)
            _set_cell_borders(vc, 'FFFFFF', '2')
            _set_cell_margin(vc, 100, 100, 160, 120)
            vp = vc.paragraphs[0]
            vr = vp.add_run(str(valor))
            vr.font.name      = 'Arial'
            vr.font.size      = Pt(10)
            vr.font.color.rgb = _hex_to_rgb(color_texto_claro)
            vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        _page_break(doc)

    @staticmethod
    def _bloque_toc(doc, estructura, tema):
        """
        TOC nativo de Word con campo TOC 1.
        - El analista abre el doc y hace clic en 'Actualizar tabla'
        - Los números de página se actualizan automáticamente
        - Requiere que _seccion_titulo use estilo Heading para que Word lo detecte
        """
        color_primario = tema.get('fondo_primario', '#1E1B4B')
        color_acento   = tema.get('acento',         '#00B4D8')

        # ── Título ───────────────────────────────────────────────
        p_titulo = doc.add_paragraph()
        p_titulo.paragraph_format.space_before = Pt(0)
        p_titulo.paragraph_format.space_after  = Pt(4)
        r = p_titulo.add_run('Tabla de Contenidos')
        r.font.name      = 'Arial'
        r.font.size      = Pt(18)
        r.font.bold      = True
        r.font.color.rgb = _hex_to_rgb(color_primario)

        _add_hr(doc, color_acento.lstrip('#'), 12)

        # ── Campo TOC nativo de Word ──────────────────────────────
        # Word detecta párrafos con estilo Heading 1 y construye el índice
        p_toc = doc.add_paragraph()

        # Instrucción del campo: TOC con nivel 1, sin hiperlinks, con líderes
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        fldChar_begin.set(qn('w:dirty'), 'true')   # fuerza re-renderizado al abrir

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-1" \\h \\z \\u '
        # \\o "1-1" → solo Heading 1
        # \\h       → hipervínculos (Ctrl+clic navega a la sección)
        # \\z       → oculta números de página en web layout
        # \\u       → usa estilos de párrafo con outlineLevel

        fldChar_sep = OxmlElement('w:fldChar')
        fldChar_sep.set(qn('w:fldCharType'), 'separate')

        # Texto placeholder que ve el usuario antes de actualizar
        r_placeholder = OxmlElement('w:r')
        t_placeholder = OxmlElement('w:t')
        t_placeholder.text = '[Haga clic en Actualizar tabla para generar el índice]'
        r_placeholder.append(t_placeholder)

        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')

        # Ensamblar el campo en el run del párrafo
        run = p_toc.add_run()
        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_sep)
        run._r.append(r_placeholder)
        run._r.append(fldChar_end)

        _page_break(doc)

    @staticmethod
    def _seccion_titulo(doc, texto, tema):
        color_primario = tema.get('fondo_primario', '#1E1B4B')
        color_acento   = tema.get('acento',         '#00B4D8')

        p = doc.add_paragraph(style='Heading 1')  # ← clave para el TOC
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)

        r = p.add_run(texto)
        r.font.name      = 'Arial'
        r.font.size      = Pt(15)
        r.font.bold      = True
        r.font.color.rgb = _hex_to_rgb(color_primario)

        _add_hr(doc, color_acento.lstrip('#'), 10)

    @staticmethod
    def _bloque_seccion_estatica(doc, subtitulo, contenido, tema):
        ReportService._seccion_titulo(doc, subtitulo, tema)

        texto = contenido or '[Completar por el analista]'
        texto = texto.replace('\r\n', '\n').replace('\r', '\n')  # normaliza CRLF/CR sueltos
        parrafos = [p.strip() for p in texto.split('\n\n') if p.strip()]

        color = tema.get('texto_oscuro', '#111827') if contenido else tema.get('borde', '#CCCCCC')

        for parrafo_texto in parrafos:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            lineas = parrafo_texto.split('\n')
            for i, linea in enumerate(lineas):
                r = p.add_run(linea)
                r.font.name      = 'Arial'
                r.font.size      = Pt(10)
                r.font.italic    = contenido is None
                r.font.color.rgb = _hex_to_rgb(color)
                if i < len(lineas) - 1:
                    r.add_break()

        doc.add_paragraph()
        
    @staticmethod
    def _bloque_seccion_vacia(doc, subtitulo, tema):
        """Sección estática — el analista la completa a mano."""
        ReportService._seccion_titulo(doc, subtitulo, tema)
        p = doc.add_paragraph()
        r = p.add_run('[Completar por el analista]')
        r.font.name      = 'Arial'
        r.font.size      = Pt(10)
        r.font.italic    = True
        r.font.color.rgb = _hex_to_rgb(tema.get('borde', '#CCCCCC'))
        doc.add_paragraph()

    # ─────────────────────────────────────────────────────────────
    # DOCX — punto de entrada público (OSINT)
    # ─────────────────────────────────────────────────────────────
    @staticmethod
    def generar_docx(data, proyecto, tema, estructura, severidades, contenido_secciones=None, base_dir=None, tipo_informe='tecnico'):
        contenido_secciones = contenido_secciones or {}
        config = ReportService.TIPOS_INFORME.get(tipo_informe, ReportService.TIPOS_INFORME['tecnico'])

        for f in data:
            imagenes_raw  = f.get('imagenes') or ''
            f['evidencias'] = [img for img in imagenes_raw.split('|') if img]

        doc = ReportService._doc_base(proyecto, tema)
        ReportService._bloque_portada(doc, proyecto, tema, tipo_informe=tipo_informe)
        ReportService._bloque_toc(doc, estructura, tema)

        secciones_render = [
            s for s in estructura
            if s['tipo'] not in ('portada', 'toc') and s['clave'] not in config.get('excluir_secciones', [])
        ]
        ultima_clave = secciones_render[-1]['clave'] if secciones_render else None

        for seccion in estructura:
            tipo      = seccion['tipo']
            clave     = seccion['clave']
            subtitulo = seccion['subtitulo']
            dinamico  = seccion['es_dinamico']

            if tipo in ('portada', 'toc'):
                continue

            if clave in config.get('excluir_secciones', []):
                continue

            if dinamico:
                if clave == 'resumen_hallazgos':
                    ReportServiceAWS._bloque_resumen(doc, data, tema, severidades)
                elif clave == 'hallazgos':
                    ReportServiceAWS._bloque_tabla_hallazgos(doc, data, tema, severidades)
                elif clave == 'detalle_hallazgos':
                    ReportServiceAWS._bloque_detalle_hallazgos(
                        doc, data, tema, severidades,
                        base_dir=base_dir,
                        incluir_evidencia=config['incluir_evidencia']
                    )
                elif clave == 'analisis_exposicion':
                    ReportServiceAWS._bloque_analisis_exposicion(doc, data, tema, severidades)
                elif clave == 'mitre_attack':
                    ReportServiceAWS._bloque_mitre_attack(doc, data, tema, base_dir)
            else:
                contenido = contenido_secciones.get(clave)
                if clave == 'anexo_clasificacion':
                    ReportServiceAWS._bloque_anexo_clasificacion(doc, contenido, tema, severidades)
                else:
                    ReportService._bloque_seccion_estatica(doc, subtitulo, contenido, tema)

            if clave != ultima_clave:
                _page_break(doc)

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output